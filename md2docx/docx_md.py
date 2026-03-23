"""
Markdown to DOCX Converter using pypandoc with Mermaid support

This module provides functionality to convert Markdown documents to DOCX format
using pandoc. It supports custom templates, LaTeX formulas, and Mermaid diagrams
rendered via Playwright.
"""

import argparse
import logging
import os
import re
import shutil
import tempfile
from pathlib import Path
from typing import Optional, Tuple
from PIL import Image

logger = logging.getLogger(__name__)


class PandocMdDocxConverter:
    """
    Converter for transforming Markdown documents to DOCX format using pandoc.

    Features:
    - LaTeX formula support (native pandoc)
    - Custom DOCX templates
    - Mermaid diagram support via Playwright
    - Table of contents generation
    - Code highlighting
    - Academic table styling with configurable styles
    """

    def __init__(
        self,
        template_path: Optional[str] = None,
        use_default_template: bool = False,
        mermaid_images_dir: str = "files/temp_mermaid_images",
        mermaid_dpi: int = 600,
        max_image_width_cm: float = 16.0,
        max_image_height_cm: float = 25.0,
        add_image_border: bool = False,
        enhance_tables: bool = True,
        table_style: str = "academic_three_line",
    ):
        """
        Initialize the DOCX converter.

        Args:
            template_path: Path to custom DOCX template file
            use_default_template: Whether to use the default template
            mermaid_images_dir: Directory for storing Mermaid diagram images
            mermaid_dpi: DPI value for Mermaid diagrams (96-600, default: 600)
            max_image_width_cm: Maximum image width in centimeters (default: 16.0cm)
            max_image_height_cm: Maximum image height in centimeters (default: 25.0cm)
            add_image_border: Whether to add border to images (default: False)
            enhance_tables: Whether to apply academic table styling (default: True)
            table_style: Table style name (default: 'academic_three_line')
        """
        self.template_path = template_path
        self.use_default_template = use_default_template
        self.mermaid_images_dir = Path(mermaid_images_dir)
        self.mermaid_dpi = max(96, min(600, mermaid_dpi))

        # Image size control parameters
        self.max_image_width_cm = max_image_width_cm
        self.max_image_height_cm = max_image_height_cm
        self.add_image_border = add_image_border

        # Table styling
        self.enhance_tables = enhance_tables
        self.table_style = table_style

        # Create files directory if it doesn't exist
        if "files" in str(self.mermaid_images_dir):
            files_dir = Path("files")
            if not files_dir.exists():
                files_dir.mkdir(parents=True, exist_ok=True)
                logger.info(f"Created directory: {files_dir}")

        # Set default template path if requested
        if use_default_template and not template_path:
            self.template_path = str(Path(__file__).parent / "templates" / "docx_template.docx")

    def convert_file_to_docx(
        self,
        input_file: str,
        output_file: str,
        process_mermaid: bool = True,
        cleanup_temp: bool = False,
    ) -> str:
        """
        Convert a Markdown file to DOCX format.

        Args:
            input_file: Path to input Markdown file
            output_file: Path to output DOCX file
            process_mermaid: Whether to convert Mermaid diagrams to images
            cleanup_temp: Whether to clean up temporary files (default: False to keep images)

        Returns:
            Path to the generated DOCX file
        """
        input_path = Path(input_file)
        output_path = Path(output_file)

        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        # Create output directory if needed
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Converting {input_file} to {output_file}")

        try:
            # Read Markdown content
            with open(input_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            # Preprocess Markdown (Mermaid diagrams)
            temp_dir = None
            if process_mermaid and self._has_mermaid_diagrams(md_content):
                logger.info("Processing Mermaid diagrams...")
                md_content, temp_dir = self._process_mermaid_diagrams(md_content, input_path.parent)

            # Create temporary Markdown file for pandoc
            with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as temp_md:
                temp_md.write(md_content)
                temp_md_path = temp_md.name

            try:
                # Convert using pandoc
                self._convert_with_pandoc(temp_md_path, str(output_path))
                logger.info(f"Successfully converted to {output_file}")

            finally:
                # Clean up temporary markdown file
                os.unlink(temp_md_path)

            # Clean up Mermaid images directory if requested
            if cleanup_temp and temp_dir and temp_dir.exists():
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")

            return str(output_file)

        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            raise

    def _has_mermaid_diagrams(self, md_content: str) -> bool:
        """Check if Markdown content contains Mermaid diagrams."""
        return bool(re.search(r"```mermaid", md_content))

    def _process_mermaid_diagrams(self, md_content: str, base_dir: Path) -> Tuple[str, Optional[Path]]:
        """
        Convert Mermaid diagrams to images using Playwright.

        Args:
            md_content: Markdown content
            base_dir: Base directory for resolving paths

        Returns:
            Tuple of (updated markdown content, temp directory path)
        """
        # Create temporary directory for images
        # Use absolute path from mermaid_images_dir to avoid path concatenation issues
        temp_dir = self.mermaid_images_dir.resolve()
        temp_dir.mkdir(parents=True, exist_ok=True)

        # Find all Mermaid code blocks
        mermaid_pattern = re.compile(r"```mermaid\n(.*?)\n```", re.DOTALL)
        mermaid_blocks = list(mermaid_pattern.finditer(md_content))

        if not mermaid_blocks:
            return md_content, None

        # Try to use Playwright to render Mermaid
        try:
            updated_content = self._render_mermaid_with_playwright(md_content, mermaid_blocks, temp_dir, base_dir)
            return updated_content, temp_dir
        except ImportError:
            logger.warning("Playwright not available. Mermaid diagrams will remain as code blocks.")
            logger.warning("Install playwright: pip install playwright && playwright install chromium")
            return md_content, temp_dir
        except Exception as e:
            logger.warning(f"Failed to render Mermaid diagrams: {e}")
            logger.warning("Mermaid diagrams will remain as code blocks.")
            return md_content, temp_dir

    def _fix_mermaid_syntax(self, content: str) -> str:
        """修复 Mermaid 语法：中文引号转英文引号"""
        content = content.replace(""", '"').replace(""", '"').replace("'", "'").replace("'", "'")
        return content

    def _render_mermaid_with_playwright(
        self,
        md_content: str,
        mermaid_blocks: list,
        temp_dir: Path,
        base_dir: Path,
    ) -> str:
        """
        Render Mermaid diagrams to PNG images using Playwright.

        Args:
            md_content: Original Markdown content
            mermaid_blocks: List of regex match objects for Mermaid blocks
            temp_dir: Directory to save images
            base_dir: Base directory for file operations

        Returns:
            Updated Markdown content with Mermaid blocks replaced by image references
        """
        from playwright.sync_api import sync_playwright

        updated_content = md_content
        image_counter = 0

        with sync_playwright() as p:
            # Launch browser
            browser = p.chromium.launch()

            # Calculate device scale factor (standard screen DPI is 96)
            scale_factor = self.mermaid_dpi / 96.0

            # Start with a reasonable base viewport
            base_width = 1600
            base_height = 1200

            # Create initial page
            page = browser.new_page(
                viewport={"width": base_width, "height": base_height},
                device_scale_factor=scale_factor,
            )

            # Process each Mermaid diagram in reverse order
            for match in reversed(mermaid_blocks):
                mermaid_code = match.group(1)
                # 修复中文引号等语法问题
                mermaid_code = self._fix_mermaid_syntax(mermaid_code)
                image_filename = f"mermaid_{image_counter}.png"
                image_path = temp_dir / image_filename

                # Create HTML page with Mermaid code
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{
                            margin: 0;
                            padding: 20px;
                            display: flex;
                            justify-content: center;
                            align-items: center;
                            background: white;
                            min-height: 100vh;
                        }}
                        .mermaid {{
                            background: white;
                            padding: 15px;
                            font-size: 16px;
                        }}
                        svg {{
                            max-width: none !important;
                            height: auto !important;
                            display: block !important;
                        }}
                    </style>
                </head>
                <body>
                    <pre class="mermaid">
{mermaid_code}
                    </pre>
                    <script type="module">
                        import mermaid from 'https://unpkg.com/mermaid@11/dist/mermaid.esm.min.mjs';

                        mermaid.initialize({{
                            startOnLoad: true,  // 自动加载渲染
                            theme: 'default',
                            securityLevel: 'loose',
                            fontSize: 16,
                            fontFamily: 'Arial, "Microsoft YaHei", sans-serif',
                            flowchart: {{
                                useMaxWidth: false,
                                htmlLabels: true,
                                curve: 'basis',
                                padding: 10
                            }},
                            sequence: {{
                                useMaxWidth: false,
                                diagramMarginX: 10,
                                diagramMarginY: 10,
                                boxMargin: 8,
                                messageMargin: 25
                            }}
                        }});
                    </script>
                </body>
                </html>
                """

                try:
                    # Create temporary HTML file
                    with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as temp_html:
                        temp_html.write(html_content)
                        temp_html_path = temp_html.name

                    try:
                        # Load page and wait for initial load
                        page.goto(f"file://{temp_html_path}", wait_until="domcontentloaded")

                        # Wait for SVG element to appear (简化等待方式)
                        try:
                            page.wait_for_selector(".mermaid svg", timeout=10000)
                        except:
                            pass  # Continue anyway

                        # Wait for rendering to complete
                        page.wait_for_timeout(2000)

                        # Get the Mermaid element and screenshot it
                        mermaid_element = page.query_selector(".mermaid")
                        if mermaid_element:
                            # Screenshot with high quality
                            mermaid_element.screenshot(
                                path=str(image_path),
                                type="png",
                                scale="device",  # Use device pixel ratio
                            )

                            # 检查并缩放图片：使用改进的尺寸控制逻辑
                            try:
                                with Image.open(image_path) as img:
                                    width, height = img.size

                                    # 转换为厘米
                                    width_cm = width * 2.54 / self.mermaid_dpi
                                    height_cm = height * 2.54 / self.mermaid_dpi

                                    logger.info(f"原始图片尺寸: {width}x{height}px ({width_cm:.2f}×{height_cm:.2f}cm)")

                                    # 使用初始化时设置的尺寸限制
                                    max_width_cm = self.max_image_width_cm
                                    max_height_cm = self.max_image_height_cm

                                    # 检查是否需要缩放
                                    needs_resize = False
                                    new_width, new_height = width, height

                                    if width_cm > max_width_cm or height_cm > max_height_cm:
                                        needs_resize = True

                                        # 计算宽高比例
                                        width_ratio = max_width_cm / width_cm if width_cm > max_width_cm else 1.0
                                        height_ratio = max_height_cm / height_cm if height_cm > max_height_cm else 1.0

                                        # 使用较小的比例以确保图片完全在限制内
                                        scale_factor = min(width_ratio, height_ratio)

                                        new_width = int(width * scale_factor)
                                        new_height = int(height * scale_factor)

                                        # 缩放图片
                                        resized_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                                        resized_img.save(image_path)

                                        new_width_cm = new_width * 2.54 / self.mermaid_dpi
                                        new_height_cm = new_height * 2.54 / self.mermaid_dpi

                                        logger.info(f"图片已缩放: {width}x{height}px ({width_cm:.1f}×{height_cm:.1f}cm) -> {new_width}x{new_height}px ({new_width_cm:.1f}×{new_height_cm:.1f}cm)")
                                    else:
                                        logger.info(f"图片尺寸合适: {width}x{height}px ({width_cm:.1f}×{height_cm:.1f}cm)，无需缩放")

                                    # 可选：添加图片边框
                                    if self.add_image_border:
                                        try:
                                            border_width = 3
                                            border_color = (200, 200, 200)  # 浅灰色
                                            bordered_img = Image.new('RGB',
                                                (new_width + border_width*2, new_height + border_width*2),
                                                border_color
                                            )
                                            bordered_img.paste(Image.open(image_path), (border_width, border_width))
                                            bordered_img.save(image_path)
                                            logger.info(f"已添加图片边框: {border_width}px")
                                        except Exception as border_err:
                                            logger.warning(f"添加边框失败: {border_err}")

                                    # 设置图片DPI信息
                                    try:
                                        img_with_dpi = Image.open(image_path)
                                        img_with_dpi.save(image_path, dpi=(self.mermaid_dpi, self.mermaid_dpi))
                                    except:
                                        pass  # DPI设置失败不影响继续处理

                            except Exception as e:
                                logger.warning(f"图片尺寸检查或缩放失败: {e}")

                            # Replace code block with image reference
                            # Convert path to string and use forward slashes for Markdown
                            image_abs_path = image_path.resolve()
                            image_path_str = str(image_abs_path).replace("\\", "/")

                            image_ref = f'\n\n![]({image_path_str})\n\n'
                            updated_content = (
                                updated_content[: match.start()] + image_ref + updated_content[match.end() :]
                            )
                            image_counter += 1
                            logger.info(f"Converted Mermaid diagram to {image_filename}")

                    finally:
                        os.unlink(temp_html_path)

                except Exception as e:
                    logger.warning(f"Failed to convert Mermaid diagram: {e}")
                    continue

            logger.info("All Mermaid diagrams processed, proceeding to conversion...")
            # 不显式关闭浏览器，让进程退出时自动清理，避免线程安全问题

        return updated_content

    def _convert_with_pandoc(self, input_path: str, output_path: str):
        """
        Convert Markdown to DOCX using pypandoc and apply table styling.

        Args:
            input_path: Path to input Markdown file
            output_path: Path to output DOCX file
        """
        try:
            import pypandoc
        except ImportError:
            raise ImportError(
                "pypandoc-binary is not installed. "
                "Please install it using: pip install pypandoc-binary"
            )

        # Build extra arguments
        extra_args = [
            "--syntax-highlighting=pygments",
        ]

        # Add template if specified
        if self.template_path:
            extra_args.append(f"--reference-doc={self.template_path}")

        logger.info(f"Converting to DOCX, extra args: {extra_args}")

        # Perform conversion
        pypandoc.convert_file(
            source_file=input_path,
            format="markdown",
            to="docx",
            outputfile=output_path,
            extra_args=extra_args,
        )

        # Apply table styling to the generated DOCX
        if self.enhance_tables:
            logger.info("Applying academic table styling to DOCX...")
            self._apply_table_styling_to_docx(output_path)

    def _apply_table_styling_to_docx(self, docx_path: str):
        """
        Apply academic table styling to the generated DOCX file using python-docx.

        Args:
            docx_path: Path to the DOCX file to style
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed. Skipping table styling.")
            logger.warning("Install with: pip install python-docx")
            return

        try:
            # Load table style configuration
            import sys
            from pathlib import Path

            # Add the templates directory to the path
            templates_dir = Path(__file__).parent / "templates"
            if str(templates_dir) not in sys.path:
                sys.path.insert(0, str(templates_dir))

            from table_styles import get_table_style
            style_config = get_table_style(self.table_style)

            logger.info(f"Applying table style: {style_config['name']}")

            doc = Document(docx_path)
            table_count = 0

            for table in doc.tables:
                table_count += 1
                logger.info(f"Processing table with {len(table.rows)} rows")

                # Get the first and last row indices
                first_row_idx = 0
                last_row_idx = len(table.rows) - 1

                for i, row in enumerate(table.rows):
                    for cell in row.cells:
                        is_header = (i == first_row_idx)
                        is_last_row = (i == last_row_idx)
                        is_odd_row = (i % 2 == 1)

                        if is_header:
                            logger.debug(f"Processing header row {i}, cell text: {cell.text[:20] if cell.text else '(empty)'}")

                        # Apply header styling
                        if is_header:
                            # Set background color
                            bg_color = style_config["header"]["background"]
                            self._set_cell_shading(cell, bg_color)

                            # Set text formatting
                            for paragraph in cell.paragraphs:
                                # Set alignment
                                align_str = style_config["header"]["align"]
                                if align_str == "center":
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif align_str == "left":
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                elif align_str == "right":
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                                # Set bold and font size (color uses default black)
                                for run in paragraph.runs:
                                    if style_config["header"]["bold"]:
                                        run.font.bold = True
                                    if "font_size" in style_config["header"]:
                                        from docx.shared import Pt
                                        run.font.size = Pt(style_config["header"]["font_size"])
                                    # Note: color is ignored, uses default black

                        # Apply row background colors
                        else:
                            if is_odd_row:
                                bg_color = style_config["rows"]["odd"]["background"]
                            else:
                                bg_color = style_config["rows"]["even"]["background"]
                            self._set_cell_shading(cell, bg_color)

                        # Apply borders based on position
                        self._apply_style_borders(cell, is_header, is_last_row, style_config)

            doc.save(docx_path)
            logger.info(f"Applied '{style_config['name']}' styling to {table_count} table(s)")

        except Exception as e:
            logger.warning(f"Failed to apply table styling: {e}")
            import traceback
            logger.warning(traceback.format_exc())

    def _set_cell_shading(self, cell, fill_color):
        """Set cell background color."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._element.get_or_add_tcPr()

        # Remove existing shading if any
        existing_shd = tc_pr.find(qn('w:shd'))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)

        # Create and add new shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_color)
        tc_pr.append(shd)

    def _set_cell_borders(self, cell, header=False):
        """Set cell borders with thicker borders for header row."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._element.get_or_add_tcPr()

        # Check if borders already exist
        existing_borders = tc_pr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            tc_pr.remove(existing_borders)

        tc_borders = OxmlElement('w:tcBorders')

        # Border size: 4 for regular, 12 for header (0.5pt vs 1.5pt)
        border_size = '12' if header else '4'

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), border_size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tc_borders.append(border)

        tc_pr.append(tc_borders)

    def _apply_style_borders(self, cell, is_header, is_last_row, style_config):
        """
        Apply borders to a cell based on style configuration.

        Args:
            cell: The cell to apply borders to
            is_header: Whether this cell is in the header row
            is_last_row: Whether this cell is in the last row
            style_config: The style configuration dictionary
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._element.get_or_add_tcPr()

        # Remove existing borders
        existing_borders = tc_pr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            tc_pr.remove(existing_borders)

        tc_borders = OxmlElement('w:tcBorders')
        borders = style_config["borders"]

        # Top border
        if is_header:
            border_config = borders["top"]
        else:
            border_config = borders.get("inner_horizontal", borders["top"])
        self._add_border(tc_borders, 'top', border_config)

        # Bottom border
        if is_header:
            border_config = borders["header_bottom"]
        elif is_last_row:
            border_config = borders["bottom"]
        else:
            border_config = borders.get("inner_horizontal", {"size": "0", "color": "000000", "style": "none"})
        self._add_border(tc_borders, 'bottom', border_config)

        # Left border
        self._add_border(tc_borders, 'left', borders["left"])

        # Right border
        self._add_border(tc_borders, 'right', borders["right"])

        tc_pr.append(tc_borders)

    def _add_border(self, tc_borders, border_name, border_config):
        """Add a single border to the borders element."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), border_config.get("style", "single"))
        border.set(qn('w:sz'), str(border_config.get("size", "4")))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_config.get("color", "000000"))
        tc_borders.append(border)


def main_helper(input_file, output_file, template=None, use_default_template=False,
                mermaid_dpi=600, no_mermaid=False, image_quality=85, max_image_width=15, cleanup_temp=False,
                max_image_width_cm=16.0, max_image_height_cm=25.0,
                add_image_border=False, enhance_tables=True, table_style="academic_three_line"):
    """Helper function for calling from main entry point."""
    import logging

    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )

    # Create converter
    converter = PandocMdDocxConverter(
        template_path=template,
        use_default_template=use_default_template,
        mermaid_dpi=mermaid_dpi,
        max_image_width_cm=max_image_width_cm,
        max_image_height_cm=max_image_height_cm,
        add_image_border=add_image_border,
        enhance_tables=enhance_tables,
        table_style=table_style,
    )

    # Perform conversion
    converter.convert_file_to_docx(
        input_file=input_file,
        output_file=output_file,
        process_mermaid=not no_mermaid,
        cleanup_temp=cleanup_temp,
    )
    if not cleanup_temp:
        print(f"[INFO] Mermaid images saved to: files/temp_mermaid_images")
    print(f"[OK] Successfully converted {input_file} to {output_file}")


def main():
    """Command-line interface for the DOCX converter."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX using pandoc with Mermaid support",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic conversion (default: three-line table style)
  python docx_md.py -i input.md -o output.docx

  # Academic three-line table style (三线表 - 中国学术标准)
  python docx_md.py -i input.md -o output.docx --table-style academic_three_line

  # Academic striped table style (条纹表 - 适合大量数据)
  python docx_md.py -i input.md -o output.docx --table-style academic_striped

  # Simple grid table style (简单网格 - 简洁清晰)
  python docx_md.py -i input.md -o output.docx --table-style simple_grid

  # Minimal table style (极简表格 - 现代简洁)
  python docx_md.py -i input.md -o output.docx --table-style minimal

  # Convert with custom image size limits
  python docx_md.py -i input.md -o output.docx --max-image-width-cm 15 --max-image-height-cm 25

  # Convert with borders on Mermaid diagrams
  python docx_md.py -i input.md -o output.docx --add-image-border

  # Use custom template
  python docx_md.py -i input.md -o output.docx -t my_template.docx

  # Disable Mermaid processing
  python docx_md.py -i input.md -o output.docx --no-mermaid

  # Convert without table styling
  python docx_md.py -i input.md -o output.docx --no-table-enhance

Table Styles:
  - academic_three_line: 三线表 (顶粗、中细、底粗，无竖线) [默认]
  - academic_striped: 条纹表 (斑马纹，完整边框，蓝灰配色)
  - simple_grid: 简单网格 (所有线条一致，浅灰表头)
  - minimal: 极简表格 (仅表头表底有边框)

Table Style Config: md2docx/templates/table_styles.py
        """,
    )

    parser.add_argument(
        "-i", "--input",
        required=True,
        help="Input Markdown file path (will also search in files/md_test_files/ if not found)",
    )

    parser.add_argument(
        "-o", "--output",
        default=None,
        help="Output DOCX file path (default: files/docx/<input_filename>.docx)",
    )

    parser.add_argument(
        "-t", "--template", help="Path to custom DOCX template file"
    )

    parser.add_argument(
        "--use-default-template",
        action="store_true",
        help="Use the default template from templates/docx_template.docx",
    )

    parser.add_argument(
        "--mermaid-images-dir",
        default="files/temp_mermaid_images",
        help="Directory for Mermaid diagram images (default: files/temp_mermaid_images)",
    )

    parser.add_argument(
        "--mermaid-dpi",
        type=int,
        default=600,
        help="DPI value for Mermaid diagrams (96-600, default: 600 - maximum quality). Higher values = clearer but larger files",
    )

    parser.add_argument(
        "--no-mermaid",
        action="store_true",
        help="Skip Mermaid diagram processing",
    )

    parser.add_argument(
        "--cleanup",
        action="store_true",
        help="Clean up temporary mermaid images after conversion (default: keep images)",
    )

    # Image size control parameters
    parser.add_argument(
        "--max-image-width-cm",
        type=float,
        default=16.0,
        help="Maximum image width in centimeters (default: 16.0)",
    )

    parser.add_argument(
        "--max-image-height-cm",
        type=float,
        default=25.0,
        help="Maximum image height in centimeters (default: 25.0)",
    )

    parser.add_argument(
        "--add-image-border",
        action="store_true",
        help="Add border to Mermaid diagram images (default: no border)",
    )

    parser.add_argument(
        "--enhance-tables",
        action="store_true",
        default=True,
        help="Apply academic table styling (default: True)",
    )

    parser.add_argument(
        "--no-table-enhance",
        action="store_true",
        help="Disable table styling enhancement",
    )

    parser.add_argument(
        "--table-style",
        type=str,
        default="academic_three_line",
        choices=["academic_three_line", "academic_striped", "simple_grid", "minimal"],
        help="Table style to apply (default: academic_three_line). Options: academic_three_line (三线表), academic_striped (条纹表), simple_grid (简单网格), minimal (极简表格)",
    )

    parser.add_argument(
        "-v", "--verbose", action="store_true", help="Enable verbose logging"
    )

    args = parser.parse_args()

    # Configure logging
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )

    # Create converter
    converter = PandocMdDocxConverter(
        template_path=args.template,
        use_default_template=args.use_default_template,
        mermaid_images_dir=args.mermaid_images_dir,
        mermaid_dpi=args.mermaid_dpi,
        max_image_width_cm=args.max_image_width_cm,
        max_image_height_cm=args.max_image_height_cm,
        add_image_border=args.add_image_border,
        enhance_tables=not args.no_table_enhance,
        table_style=args.table_style,
    )

    try:
        # Resolve input file path
        input_file = args.input
        input_path = Path(input_file)

        # If not found, try looking in default test files directory
        if not input_path.exists():
            test_files_dir = Path("files/md_test_files")
            test_file_path = test_files_dir / input_file
            if test_file_path.exists():
                input_file = str(test_file_path)
                logger.info(f"Found input file in test files directory: {input_file}")
            else:
                raise FileNotFoundError(f"Input file not found: {input_file}\nAlso searched in: {test_files_dir}")

        # Generate output path if not specified
        output_file = args.output
        if not output_file:
            input_path = Path(input_file)
            output_dir = Path("files/docx")
            output_dir.mkdir(parents=True, exist_ok=True)
            output_file = str(output_dir / f"{input_path.stem}.docx")
            logger.info(f"Output file not specified, using default: {output_file}")

        # Perform conversion
        converter.convert_file_to_docx(
            input_file=input_file,
            output_file=output_file,
            process_mermaid=not args.no_mermaid,
            cleanup_temp=args.cleanup,
        )
        if not args.cleanup:
            print(f"[INFO] Mermaid images saved to: {converter.mermaid_images_dir}")
        print(f"[OK] Successfully converted {input_file} to {output_file}")

    except Exception as e:
        print(f"[ERROR] Conversion failed: {e}")
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
